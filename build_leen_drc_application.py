from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.style import WD_STYLE_TYPE

OUT = Path('/home/openclaw/.openclaw/workspace/nuke-watch/deliverables/Leen_Alhamami_DRC_Application')
OUT.mkdir(parents=True, exist_ok=True)

NAVY = '17324D'
TEAL = '2B6F75'
SLATE = '4B5563'
DARK = '1F2937'
MID = '667085'
LIGHT = 'D7E0E8'
PALE = 'F4F7F9'
WHITE = 'FFFFFF'

FONT = 'Aptos'
FALLBACK = 'Arial'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is not None:
        tblPr.remove(borders)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
        '<w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def set_run_font(run, size=None, bold=None, italic=None, color=None, name=FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def keep_table_row(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Page ')
    set_run_font(run, 8, color=MID)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = paragraph.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    set_run_font(run2, 8, color=MID)


def configure_document(doc, footer_label):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.35)
    sec.bottom_margin = Cm(1.35)
    sec.left_margin = Cm(1.55)
    sec.right_margin = Cm(1.55)
    sec.header_distance = Cm(0.5)
    sec.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    normal.font.size = Pt(10.1)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    pf = normal.paragraph_format
    pf.space_after = Pt(3.2)
    pf.line_spacing = 1.08

    for style_name in ['Title', 'Subtitle', 'Heading 1', 'Heading 2']:
        styles[style_name].font.name = FONT
        styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    if 'CV Section' not in [s.name for s in styles]:
        st = styles.add_style('CV Section', WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles['CV Section']
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.font.size = Pt(10.2)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(NAVY)
    st.font.all_caps = True
    st.paragraph_format.space_before = Pt(8)
    st.paragraph_format.space_after = Pt(3)
    st.paragraph_format.keep_with_next = True

    if 'CV Bullet' not in [s.name for s in styles]:
        bst = styles.add_style('CV Bullet', WD_STYLE_TYPE.PARAGRAPH)
    else:
        bst = styles['CV Bullet']
    bst.font.name = FONT
    bst._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    bst.font.size = Pt(9.7)
    bst.font.color.rgb = RGBColor.from_string(DARK)
    bst.paragraph_format.left_indent = Cm(0.42)
    bst.paragraph_format.first_line_indent = Cm(-0.28)
    bst.paragraph_format.space_after = Pt(1.7)
    bst.paragraph_format.line_spacing = 1.06

    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="5" w:space="5" w:color="{LIGHT}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    r = p.add_run(footer_label)
    set_run_font(r, 8, color=MID)
    r2 = p.add_run('                                                         ')
    set_run_font(r2, 8, color=MID)
    add_page_field(p)

    # Metadata
    doc.core_properties.author = 'Leen Alhamami'
    doc.core_properties.subject = footer_label
    doc.core_properties.keywords = 'Protection Monitoring Assistant, Danish Refugee Council, Rural Damascus'


def add_bottom_rule(paragraph, color=TEAL, size=8, space=3):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_cv_header(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('LEEN ALHAMAMI')
    set_run_font(r, 24, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('COMMUNITY-FOCUSED PROFESSIONAL')
    set_run_font(r, 10.2, bold=True, color=TEAL)
    r.font.letter_spacing = Pt(0.4) if hasattr(r.font, 'letter_spacing') else None
    add_bottom_rule(p, color=TEAL, size=9, space=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(5)
    parts = [
        ('Damascus, Syria', False),
        ('  |  ', False),
        ('+963 935 157 005', False),
        ('  |  ', False),
        ('leenalhamami65@gmail.com', False),
    ]
    for text, bold in parts:
        r = p.add_run(text)
        set_run_font(r, 9.2, bold=bold, color=SLATE)


def add_section(doc, title):
    p = doc.add_paragraph(style='CV Section')
    p.add_run(title)
    add_bottom_rule(p, color=LIGHT, size=5, space=2)
    return p


def add_bullet(doc, text, after=1.7):
    p = doc.add_paragraph(style='CV Bullet')
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run('•  ')
    set_run_font(r, 9.7, bold=True, color=TEAL)
    r = p.add_run(text)
    set_run_font(r, 9.7, color=DARK)
    return p


def add_job(doc, title, organization, dates, bullets):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    remove_table_borders(t)
    t.columns[0].width = Cm(13.5)
    t.columns[1].width = Cm(4.0)
    left, right = t.rows[0].cells
    set_cell_margins(left, 0, 0, 0, 0)
    set_cell_margins(right, 0, 0, 0, 0)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, 10.0, bold=True, color=DARK)
    r = p.add_run(f'  |  {organization}')
    set_run_font(r, 9.9, bold=True, color=TEAL)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(dates)
    set_run_font(r, 9.1, bold=True, color=MID)
    keep_table_row(t.rows[0])
    for b in bullets:
        add_bullet(doc, b, after=1.25)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0.5)
    spacer.paragraph_format.line_spacing = 0.35
    return t


def add_competency_box(doc, items):
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.3)] * 4
    for i, item in enumerate(items):
        row, col = divmod(i, 4)
        cell = table.cell(row, col)
        cell.width = widths[col]
        set_cell_margins(cell, top=70, start=85, bottom=70, end=85)
        set_cell_shading(cell, PALE if (row + col) % 2 == 0 else WHITE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item)
        set_run_font(r, 8.8, bold=True, color=SLATE)
    # subtle borders
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="3" w:color="{LIGHT}"/>'
        f'<w:left w:val="single" w:sz="3" w:color="{LIGHT}"/>'
        f'<w:bottom w:val="single" w:sz="3" w:color="{LIGHT}"/>'
        f'<w:right w:val="single" w:sz="3" w:color="{LIGHT}"/>'
        f'<w:insideH w:val="single" w:sz="3" w:color="{LIGHT}"/>'
        f'<w:insideV w:val="single" w:sz="3" w:color="{LIGHT}"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    for row in table.rows:
        keep_table_row(row)
    return table


def build_cv():
    doc = Document()
    configure_document(doc, 'Leen Alhamami | Curriculum Vitae')
    add_cv_header(doc)

    add_section(doc, 'Professional Profile')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    text = (
        'Early-career professional with experience in education, supervision, technical support, and '
        'UNDP-supported community initiatives. Native Arabic speaker with excellent English and strong '
        'communication, organization, problem-solving, and teamwork skills. Motivated to apply careful '
        'documentation, respectful community engagement, and a commitment to dignity, confidentiality, '
        'non-discrimination, and continuous learning in humanitarian protection work.'
    )
    r = p.add_run(text)
    set_run_font(r, 9.8, color=DARK)

    add_section(doc, 'Core Competencies')
    add_competency_box(doc, [
        'Respectful communication',
        'Community-focused support',
        'Team coordination',
        'Documentation & follow-up',
        'Problem-solving',
        'Time & task management',
        'Microsoft Office',
        'Arabic–English communication',
    ])

    add_section(doc, 'Professional Experience')
    add_job(doc, 'Supervisor', 'Golden Steps School', '2025–Present', [
        'Support daily supervision and coordination in a school setting, helping maintain organized and respectful interactions with students and colleagues.',
        'Manage competing responsibilities, respond calmly to routine issues, and follow up to keep activities on schedule.',
        'Communicate instructions clearly and contribute to a cooperative working environment.',
    ])
    add_job(doc, 'Technical Support', 'Learning Go Company', '2026–Present', [
        'Provide user-focused technical support by listening carefully, identifying issues, explaining solutions clearly, and following up on requests.',
        'Coordinate with colleagues when issues require additional support while maintaining attention to detail.',
    ])
    add_job(doc, 'English Language Teacher', 'Lingua Zone', '2025–Present', [
        'Plan and deliver English lessons adapted to learners’ levels and needs.',
        'Monitor learner progress, provide constructive feedback, and maintain organized teaching materials and schedules.',
    ])
    add_job(doc, 'English Teacher', 'Golden Steps School', '2024–2026', [
        'Delivered structured lessons, encouraged participation, and managed classroom activities in a respectful learning environment.',
        'Worked with colleagues to support student engagement and consistent learning outcomes.',
    ])
    add_job(doc, 'Private English Tutor', 'Self-employed', '2024–Present', [
        'Prepare individualized lessons, adapt explanations to different learning needs, and manage appointments and deadlines independently.',
    ])

    add_section(doc, 'Community & Development Experience')
    add_job(doc, 'Volunteer', 'UNDP REEF UP Project', '2021', [
        'Participated as a volunteer in a UNDP-supported community initiative.',
    ])
    add_job(doc, 'Trainee', 'UNDP Souk Sarouja Project', '2022', [
        'Completed project-based training through a UNDP-supported initiative.',
    ])

    add_section(doc, 'Education')
    add_job(doc, 'Bachelor’s Degree in English Literature', 'Damascus University', '2019–2024', [])

    add_section(doc, 'Training & Certificates')
    add_bullet(doc, 'TESOL Certificate — University of Arizona via Coursera (in progress)')
    add_bullet(doc, 'English Grammar — Coursera')
    add_bullet(doc, 'How to Create an Online Course — Coursera')
    add_bullet(doc, 'Voice-over Training — International Center for Training and Media Skills Development (2024)')

    add_section(doc, 'Languages')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    for i, (lang, level) in enumerate([
        ('Arabic', 'Native'), ('English', 'Excellent'), ('Spanish', 'Beginner')
    ]):
        if i:
            r = p.add_run('     |     ')
            set_run_font(r, 9.3, color=LIGHT)
        r = p.add_run(f'{lang}: ')
        set_run_font(r, 9.4, bold=True, color=TEAL)
        r = p.add_run(level)
        set_run_font(r, 9.4, color=DARK)

    out = OUT / 'Leen_Alhamami_CV_DRC.docx'
    doc.save(str(out))
    return out


def add_letter_header(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('LEEN ALHAMAMI')
    set_run_font(r, 21, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    for i, text in enumerate(['Damascus, Syria', '+963 935 157 005', 'leenalhamami65@gmail.com']):
        if i:
            r = p.add_run('  |  ')
            set_run_font(r, 9.0, color=LIGHT)
        r = p.add_run(text)
        set_run_font(r, 9.2, color=SLATE)
    add_bottom_rule(p, color=TEAL, size=9, space=4)


def add_letter_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.17
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, 10.2, bold=True, color=DARK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, 10.2, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, 10.2, color=DARK)
    return p


def build_cover_letter():
    doc = Document()
    configure_document(doc, 'Leen Alhamami | Cover Letter | DRC Job ID 176474')
    add_letter_header(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('29 July 2026')
    set_run_font(r, 9.6, color=MID)

    for line, bold in [
        ('Hiring Committee', True),
        ('Danish Refugee Council', True),
        ('Rural Damascus, Syria', False),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0.8)
        r = p.add_run(line)
        set_run_font(r, 9.8, bold=bold, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(11)
    r = p.add_run('RE: Protection Monitoring Assistant (Re-advertised) — Job ID 176474')
    set_run_font(r, 10.2, bold=True, color=NAVY)
    add_bottom_rule(p, color=LIGHT, size=5, space=3)

    add_letter_paragraph(doc, 'Dear Hiring Committee,')

    add_letter_paragraph(doc,
        'I am applying for the Protection Monitoring Assistant position with the Danish Refugee Council in Rural Damascus. '
        'My background combines education, school supervision, technical support, and participation in UNDP-supported '
        'community initiatives. These experiences have strengthened my ability to communicate respectfully, organize multiple '
        'responsibilities, solve problems, and support people with patience and professionalism.')

    add_letter_paragraph(doc,
        'Through my work as an English teacher, tutor, and supervisor, I have learned to listen carefully, adapt communication '
        'to different needs, maintain structured schedules and records, and work constructively with students and colleagues. '
        'My technical-support role has further developed my attention to detail, clear explanation, follow-up, and coordination '
        'when an issue requires additional support. As a native Arabic speaker with excellent English, I can communicate '
        'effectively with community members and colleagues in both languages.')

    add_letter_paragraph(doc,
        'I understand that protection monitoring requires sensitivity, confidentiality, accurate information handling, '
        'non-discrimination, and consistent respect for the dignity and safety of affected people. I am motivated to build on my '
        'transferable experience through DRC training and supervision, particularly in the Protection Analytical Framework, '
        'Protection Information Management, KoBo-based data collection, safe referrals, service mapping, and psychological first aid. '
        'I am also prepared to support field visits, community engagement, work planning, reporting, and team coordination across Rural Damascus.')

    add_letter_paragraph(doc,
        'DRC’s values of humanity, respect, independence and neutrality, participation, honesty, and transparency strongly '
        'resonate with me. I would bring a collaborative attitude, willingness to learn, careful follow-through, and a genuine '
        'commitment to supporting people affected by conflict. My approach also aligns with DRC’s core competencies: striving '
        'for excellence, collaborating, taking responsibility, communicating openly, and demonstrating integrity.')

    add_letter_paragraph(doc,
        'Thank you for considering my application. I would welcome the opportunity to discuss how my communication skills, '
        'community-focused experience, and motivation to develop in humanitarian protection can contribute to DRC’s work in Rural Damascus.')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('Sincerely,')
    set_run_font(r, 10.2, color=DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('Leen Alhamami')
    set_run_font(r, 10.5, bold=True, color=NAVY)

    out = OUT / 'Leen_Alhamami_Cover_Letter_DRC.docx'
    doc.save(str(out))
    return out


if __name__ == '__main__':
    cv = build_cv()
    cl = build_cover_letter()
    print(cv)
    print(cl)
